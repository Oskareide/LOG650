#!/usr/bin/env python3
"""
Bygger LOG650 - BM - Final.docx:
  1. Beholder forsidemaler (egenerklæring, personvern, publiseringsavtale) fra malen
  2. Bruker LOG650 - BM.docx som kilde for alt faglig innhold
  3. Anvender steg-2-rettelser (tabelltitler, ligning-nummerering, Anthropic-ref, Antagelser)
  4. Anvender steg-3-rettelser (språkstil: upersonlig form, ingen fyllord)
"""
import copy, re, subprocess
from docx import Document
from docx.oxml import OxmlElement

TMPL  = '/Users/oskareide/Desktop/Mal prosjekt LOG650 v2.docx'
BM    = '/Users/oskareide/Desktop/LOG650 - BM.docx'
OUT   = '/Users/oskareide/Desktop/LOG650 - BM - Final.docx'
NS    = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── load ──────────────────────────────────────────────────────────────────────

tmpl = Document(TMPL)
bm   = Document(BM)

tmpl_body     = tmpl.element.body
bm_body       = bm.element.body
tmpl_para_map = {id(p._element): p for p in tmpl.paragraphs}
bm_para_map   = {id(p._element): p for p in bm.paragraphs}


# ── 1. Strip template content, keep only front matter (children 0–79) ─────────

print("1. Behold kun forsidemateriell fra malen ...")
tmpl_children = list(tmpl_body)
# Remove children from index 80 onward (Sammendrag placeholder, TOC, chapters)
# Keep index 0-79: title page, egenerklæring, personvern, publiseringsavtale
for child in tmpl_children[80:]:
    tag = child.tag.split('}')[-1]
    if tag != 'sectPr':
        tmpl_body.remove(child)
print(f"   Template body children after trim: {len(list(tmpl_body))}")


# ── 2. Append all BM.docx body children ───────────────────────────────────────

print("2. Kopierer innhold fra LOG650 - BM.docx ...")
count = 0
for child in list(bm_body):
    tag = child.tag.split('}')[-1]
    if tag != 'sectPr':
        tmpl_body.append(copy.deepcopy(child))
        count += 1
print(f"   {count} elementer lagt til.\n")


# ── helpers ───────────────────────────────────────────────────────────────────

def get_text(elem):
    return ''.join(t.text or '' for t in elem.iter(f'{{{NS}}}t'))

def body_children():
    return list(tmpl_body)

def para_style_name(elem):
    # Use a fresh paragraph lookup after appending
    pStyle = elem.find(f'.//{{{NS}}}pStyle')
    val = pStyle.get(f'{{{NS}}}val') if pStyle is not None else ''
    # Map Norwegian style IDs to level
    mapping = {'Overskrift1': 'Heading 1', 'Overskrift2': 'Heading 2',
               'Overskrift3': 'Heading 3'}
    return mapping.get(val, val)

def make_para(text, style_val='Normal', bold=False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(f'{{{NS}}}val', style_val)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_h2(text):
    return make_para(text, style_val='Overskrift2')

def is_h1(elem):
    pStyle = elem.find(f'.//{{{NS}}}pStyle')
    if pStyle is None:
        return False
    return pStyle.get(f'{{{NS}}}val') in ('Overskrift1',)

def is_h2(elem):
    pStyle = elem.find(f'.//{{{NS}}}pStyle')
    if pStyle is None:
        return False
    return pStyle.get(f'{{{NS}}}val') in ('Overskrift2',)


# ── 3. Steg 2a: Flytt tabelltitler fra over til under ─────────────────────────

print("3. Steg 2a – Flytter tabelltitler ...")
children = body_children()
moves = []
for i, child in enumerate(children):
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        text = get_text(child).strip()
        if text.startswith('Tabell') and i + 1 < len(children):
            nxt = children[i + 1]
            if nxt.tag.split('}')[-1] == 'tbl':
                moves.append((child, nxt))

for title_para, table in moves:
    print(f"   Flytter: '{get_text(title_para).strip()[:60]}'")
    tmpl_body.remove(title_para)
    table.addnext(copy.deepcopy(title_para))

print(f"   {len(moves)} tabelltitler flyttet.\n")


# ── 4. Steg 2b: Nummerér ligninger i Modellering ──────────────────────────────

print("4. Steg 2b – Nummererer ligninger ...")

# Build fresh para map (paragraphs now include BM content)
fresh_para_map = {id(p._element): p for p in tmpl.paragraphs}

in_modell = False
eq_num = 1

for i, child in enumerate(body_children()):
    tag = child.tag.split('}')[-1]
    if tag != 'p':
        continue

    para = fresh_para_map.get(id(child))
    if para is None:
        continue

    if para.style.name == 'Heading 1' and 'Modellering' in para.text.strip():
        in_modell = True
        continue
    if para.style.name == 'Heading 1' and in_modell:
        break

    if not in_modell:
        continue

    text = para.text.strip()

    if 'ln(d_t)' in text and eq_num == 1:
        runs = child.findall(f'{{{NS}}}r')
        if runs:
            last_t = runs[-1].find(f'{{{NS}}}t')
            if last_t is not None:
                last_t.text = (last_t.text or '').rstrip() + f'    (5.1)'
        print(f"   (5.1) Log-differensiering")
        eq_num += 1

    elif 'Φ_P' in text and eq_num == 2:
        runs = child.findall(f'{{{NS}}}r')
        if runs:
            last_t = runs[-1].find(f'{{{NS}}}t')
            if last_t is not None:
                last_t.text = (last_t.text or '').rstrip() + f'    (5.2)'
        print(f"   (5.2) SARIMA-ligning")
        eq_num += 1

    elif 'Sikkerhetslager' in text and 'SS' in text and eq_num == 3:
        runs = child.findall(f'{{{NS}}}r')
        if len(runs) >= 3:
            t1 = ''.join(t.text or '' for t in runs[0].iter(f'{{{NS}}}t')).strip()
            t2 = ''.join(t.text or '' for t in runs[1].iter(f'{{{NS}}}t')).strip()
            t3 = ''.join(t.text or '' for t in runs[2].iter(f'{{{NS}}}t')).strip()
            p1 = make_para(f'{t1}    (5.3)', 'Normal')
            p2 = make_para(f'{t2}    (5.4)', 'Normal')
            p3 = make_para(f'{t3}    (5.5)', 'Normal')
            child.addprevious(p1)
            p1.addnext(p2)
            p2.addnext(p3)
            tmpl_body.remove(child)
            print(f"   (5.3) SS, (5.4) ROP, (5.5) EOQ")
            eq_num += 3

print(f"   {eq_num - 1} ligninger nummerert.\n")


# ── 5. Steg 2c: Legg til Anthropic-referanse i Referanser ─────────────────────

print("5. Steg 2c – Anthropic-referanse ...")

anthropic_ref = (
    "Anthropic. (2025). Claude Code [Kodingsassistent]. https://claude.ai/code"
)

fresh_para_map = {id(p._element): p for p in tmpl.paragraphs}
in_ref = False
insert_pos = None

for child in body_children():
    tag = child.tag.split('}')[-1]
    if tag != 'p':
        continue
    para = fresh_para_map.get(id(child))
    if not para:
        continue
    if para.style.name == 'Heading 1' and para.text.strip() == 'Referanser':
        in_ref = True
        continue
    if para.style.name == 'Heading 1' and in_ref:
        break
    if in_ref and para.text.strip().startswith('Carbonneau') and insert_pos is None:
        insert_pos = child
        break

new_ref = make_para(anthropic_ref, 'Normal')

if insert_pos is not None:
    insert_pos.addprevious(new_ref)
    print(f"   Lagt til Anthropic (2025) før Carbonneau.\n")
else:
    # Fallback
    for child in body_children():
        para = fresh_para_map.get(id(child))
        if para and para.style.name == 'Heading 1' and para.text.strip() == 'Referanser':
            child.addnext(new_ref)
            print("   Lagt til Anthropic (2025) i Referanser (slutt).\n")
            break


# ── 6. Steg 2d: Legg til Antagelser-underavsnitt i Innledning ────────────────

print("6. Steg 2d – Antagelser-underavsnitt ...")

antagelser_tekst = [
    ("Antagelser", "heading"),
    ("Analysen bygger på følgende forutsetninger:", "normal"),
    ("1. Normalfordelt etterspørsel: Sikkerhetslageret beregnes med SS = z·σ_d·√L der z = 1,645 (95 % servicegrad). Dette forutsetter at etterspørselen er tilnærmet normalfordelt.", "normal"),
    ("2. Deterministisk og konstant ledetid: Ledetiden settes til L = 1 uke og antas å være kjent og konstant. Variasjon i ledetid er ikke modellert.", "normal"),
    ("3. Stabile kostnadsparametere: Ordrekostnaden K = kr 500 og lagerkostnadsandelen h = 25 % av vareprisen per enhet per år antas å være konstante.", "normal"),
    ("4. Stasjonære residualer etter transformasjon: Etter log-transformasjon og differensiering (d=1) antas tidsseriene å ha uavhengige residualer, som er en forutsetning for SARIMA-modellen.", "normal"),
    ("5. Sesongperiode på 13 uker: Det antas at etterspørselsmønsteret gjentar seg med en periode på s = 13 uker, tilsvarende ett kvartal.", "normal"),
]

fresh_para_map = {id(p._element): p for p in tmpl.paragraphs}
in_innl = False
avgrensninger_end = None

for child in body_children():
    tag = child.tag.split('}')[-1]
    if tag not in ('p', 'tbl'):
        continue
    if tag == 'tbl':
        if in_innl and avgrensninger_end is not None:
            avgrensninger_end = child
        continue

    para = fresh_para_map.get(id(child))
    if not para:
        continue

    if para.style.name == 'Heading 1' and para.text.strip() == '1 Introduksjon':
        in_innl = True
        continue
    if para.style.name == 'Heading 1' and in_innl:
        break

    if in_innl:
        if para.style.name == 'Heading 2' and 'Avgrensninger' in para.text:
            avgrensninger_end = child
        elif avgrensninger_end is not None and para.style.name == 'Heading 2':
            break
        elif avgrensninger_end is not None:
            avgrensninger_end = child

new_elems = []
for text, kind in antagelser_tekst:
    new_elems.append(make_h2(text) if kind == "heading" else make_para(text, 'Normal'))

if avgrensninger_end is not None:
    ref = avgrensninger_end
    for elem in new_elems:
        ref.addnext(elem)
        ref = elem
    print("   Antagelser lagt til etter Avgrensninger.\n")
else:
    print("   ADVARSEL: Fant ikke Avgrensninger-seksjonen.\n")


# ── 7. Steg 3: Språkrettelser ─────────────────────────────────────────────────

print("7. Steg 3 – Språkrettelser ...")

fixes = {
    # Sammendrag
    'Denne oppgaven undersøker':
        'Oppgaven undersøker',
    'ved gjentatte anledninger opplevd':
        'gjentatte ganger opplevd',

    # Innledning
    'Midt i denne balansegangen ligger etterspørselsprognoser':
        'I denne avveiningen er etterspørselsprognoser',
    'I takt med at tilgangen på data øker og kunstig intelligens blir mer tilgjengelig, åpner det seg nye muligheter for å automatisere og forbedre disse beslutningsprosessene.':
        'Ettersom tilgangen på data øker og kunstig intelligens blir mer tilgjengelig, skapes nye muligheter for å automatisere og forbedre disse beslutningsprosessene.',
    'driver salg av byggevarer, men kun til profesjonelle aktører':
        'driver salg av byggevarer og retter seg utelukkende mot profesjonelle aktører',

    # Teori og litteratur
    'Dette kapittelet presenterer det teoretiske grunnlaget for prosjektet. Kapittelet er organisert rundt fire sentrale temaer: lagerstyring, etterspørselsprognoser, kunstig intelligens i logistikk og KI-støttede innkjøpssystemer. Til sammen danner disse teoriene rammeverket for modellering, analyse og diskusjon i de påfølgende kapitlene.':
        'Kapittelet gjennomgår sentral litteratur som danner det teoretiske rammeverket for prosjektet. Litteraturen er organisert rundt fire temaer: lagerstyring, etterspørselsprognoser, kunstig intelligens i logistikk og KI-støttede innkjøpssystemer. Til sammen danner disse bidragene grunnlaget for modellering, analyse og diskusjon i de påfølgende kapitlene.',
    'i all lagerstyringsteorien':
        'i lagerstyringsteorien',
    'slik tilfellet er i varehandel med ERP-systemer':
        'noe som er tilfelle i varehandel der ERP-systemer gir tilgang til historiske salgsdata',
    'og er en sentral metode i dette prosjektet':
        'og er særlig relevant for sesongpreget salgsdata',
    'Likevel er det viktig å nyansere bildet: maskinlæringsmodeller':
        'Et viktig forbehold er likevel at maskinlæringsmodeller',
    'kan overprøve (overfit)':
        'kan overtilpasses (overfit)',
    'Fosso Wamba mfl., 2017':
        'Fosso Wamba mfl., 2015',
    'Fosso Wamba mfl. (2017)':
        'Fosso Wamba mfl. (2015)',

    # Casebeskrivelse
    'Dette kapittelet beskriver Byggmakker Gravdal som case for prosjektet. Formålet er å gi leseren den kontekstuelle informasjonen som er nødvendig for å forstå rammene for analysen.':
        'Kapittelet beskriver Byggmakker Gravdal som caseenhet og gir den kontekstuelle informasjonen som er nødvendig for å forstå rammene for analysen.',
    'For dette prosjektet avgrenses analysen til fire utvalgte produkter':
        'Analysen avgrenses til fire utvalgte produkter',
    'Bedriften har ved flere anledninger opplevd':
        'Bedriften har gjentatte ganger opplevd',
    'Utfordringer knyttet til dagens system inkluderer at':
        'Sentrale utfordringer ved dagens system er at',
    # NOTE: 'diskuteres nærmere i kapittel 7' kept as-is (correct in original 8-chapter structure)

    # Data og metode
    'Dette kapittelet beskriver hvordan prosjektet er gjennomført hvilke data som er brukt, hvordan de er samlet inn og bearbeidet, og hvilke analysemetoder som er valgt.':
        'Kapittelet redegjør for datagrunnlag, databehandling og analysemetoder.',
    'Tilnærmingen er deduktiv i den forstand at etablerte teorier om etterspørselsprognoser og lagerstyring danner grunnlaget for valg av metoder, og at disse metodene deretter testes mot empiriske data fra bedriften.':
        'Tilnærmingen er deduktiv: etablerte teorier om etterspørselsprognoser og lagerstyring danner grunnlaget for metodevalg, og metodene testes mot empiriske data fra bedriften.',
    'Analyseprosessen følger en femstegsstruktur i tråd med kompendiet:':
        'Analyseprosessen følger en femstegsstruktur:',
    'Datagrunnlaget for prosjektet er historiske salgsdata':
        'Datagrunnlaget er historiske salgsdata',
    'Fullstendig kode er tilgjengelig i prosjektets GitHub-repository under mappen 005 report.':
        'Fullstendig kode er tilgjengelig i GitHub-repositoriet vedlagt oppgaven.',
    'Analysene er utviklet med støtte fra Claude Code (Anthropic)':
        'Analysene er gjennomført med støtte fra Claude Code (Anthropic, 2025).',

    # Modellering
    'Dette kapittelet presenterer den matematiske og analytiske modellen som ligger til grunn for prosjektets prognose og lagerstyringssystem.':
        'Kapittelet presenterer den matematiske og analytiske modellen som ligger til grunn for prognose- og lagerstyringssystemet.',
    'i tråd med fremgangsmåten beskrevet i kompendiet':
        '(Hyndman & Athanasopoulos, 2021)',
    # NOTE: 'presenteres i kapittel 6' kept as-is (correct: Analyse og resultater = kap 6)

    # Analyse og resultater – Kilde: strip handled separately below
    'Retningen er tydelig:':
        'Resultatene viser en tydelig tendens:',
    'Generert med Python og Claude Code (KI-støttet).':
        'Generert med Python og Claude Code (Anthropic, 2025).',
    'Generert med Python og Claude Code (KI-støttet)':
        'Generert med Python og Claude Code (Anthropic, 2025)',

    # Diskusjon
    'Dette kapittelet tolker og vurderer resultatene fra analysen i lys av problemstillingen, det teoretiske rammeverket og den praktiske konteksten hos Byggmakker Gravdal.':
        'Kapittelet tolker og vurderer resultatene fra analysen i lys av problemstillingen, det teoretiske rammeverket og den praktiske konteksten hos Byggmakker Gravdal.',
    'Resultatene fra kapittel 6 viser at både SARIMA og Gradient Boosting gir bedre prognoseakkuratesse':
        'Analyseresultatene viser at både SARIMA og Gradient Boosting gir bedre prognoseakkuratesse',
    'Gradient Boosting er avhengig av at treningsdataene inneholder representativ variasjon for den perioden modellen skal predikere med kun én fullstendig sesongssyklus i treningsdataene kan modellen ha begrenset generaliseringsevne for vinterperioden.':
        'Gradient Boosting er avhengig av at treningsdataene inneholder representativ variasjon for den perioden modellen skal predikere. Med kun én fullstendig sesongssyklus i treningsdataene kan modellen ha begrenset generaliseringsevne for vinterperioden.',
    'Det er også verdt å merke seg at maskinlæringsmodellen krever':
        'Maskinlæringsmodellen krever',
    # NOTE: 'lagerstyringsstørrelsene i kapittel 6.6' kept as-is (correct in original)
    'er en faglig begrunnelse, ikke en absolutt sannhet':
        'er faglig begrunnet, men ingen fastsatt norm',
    'Det er likevel viktig å ikke overdrive konklusjonene fra simuleringen.':
        'Likevel bør konklusjonene fra simuleringen ikke overdrives.',
    'Dette prosjektet demonstrerer metoden på fire utvalgte produkter som et proof of concept.':
        'Prosjektet demonstrerer metoden på fire utvalgte produkter som et proof of concept.',

    # Konklusjon
    'Det KI-støttede innkjøpssystemet som er utviklet kombinerer':
        'Det KI-støttede innkjøpssystemet kombinerer',
    'datadrevet innkjøpspraksis ett som kan implementeres gradvis, evalueres løpende og bygges videre på etter hvert som erfaring og datagrunnlag vokser.':
        'datadrevet innkjøpspraksis – et system som kan implementeres gradvis, evalueres løpende og bygges videre på etter hvert som erfaring og datagrunnlag vokser.',
    'det krever først og fremst':
        'det krever primært',
    'Selve arbeidsprosessen med dette prosjektet har gitt en uventet, men konkret illustrasjon av nettopp dette.':
        'Arbeidsprosessen har gitt en konkret illustrasjon av dette.',
    'Å se en slik agent analysere salgsdata, identifisere mønstre, generere Python-kode og strukturere faglig tekst i sanntid er ikke bare en effektiv arbeidsmåte:':
        'Agentens evne til å analysere salgsdata, identifisere mønstre, generere Python-kode og strukturere faglig tekst i sanntid er ikke bare en effektiv arbeidsmåte:',
    'lavere enn mange antar':
        'lavere enn ofte antatt',
    'den kanskje mest verdifulle innsikten ikke alltid kommer frem':
        'verdifull innsikt ikke alltid fremkommer',
    'dette prosjektet argumenterer for å ta i bruk i innkjøpssystemer':
        'prosjektet argumenterer for å ta i bruk i innkjøpssystemer',
}

changes = 0
to_delete = []

for p in tmpl.paragraphs:
    t = p.text

    # Remove Script: file references (own paragraphs)
    if t.strip().startswith('Script:') and '004_scripts' in t:
        to_delete.append(p._element)
        changes += 1
        print(f"   Sletter Script-ref: {t[:60]}")
        continue

    # Remove standalone Kilde: paragraphs
    if t.strip().startswith('Kilde:') and ('.py' in t or '.csv' in t):
        to_delete.append(p._element)
        changes += 1
        print(f"   Sletter Kilde-ref: {t[:60]}")
        continue

    # Strip inline Kilde: from middle of paragraph text
    if 'Kilde:' in t and ('.py' in t or '.csv' in t):
        new_t = re.sub(r'\s*Kilde:.*$', '', t).strip()
        if not new_t.endswith('.'):
            new_t += '.'
        full = ''.join(r.text for r in p.runs)
        new_full = re.sub(r'\s*Kilde:.*$', '', full).strip()
        if not new_full.endswith('.'):
            new_full += '.'
        if p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ''
        changes += 1
        print(f"   Strippet Kilde fra: {new_full[:70]}")
        continue

    # Apply text fixes
    full = ''.join(r.text for r in p.runs)
    new_full = full
    for old, new in fixes.items():
        if old in new_full:
            new_full = new_full.replace(old, new)
            print(f"   Rettet: {old[:65]!r}")
    if new_full != full:
        if p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ''
        changes += 1

for elem in to_delete:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

print(f"\n   Totalt {changes} språkrettelser.\n")


# ── 8. Lagre ──────────────────────────────────────────────────────────────────

tmpl.save(OUT)
subprocess.run(['xattr', '-d', 'com.apple.quarantine', OUT],
               capture_output=True)
print(f"Lagret → {OUT}")

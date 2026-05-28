#!/usr/bin/env python3
"""
Runde 2 – rettelser i LOG650 - nyeste - ren.docx:
  1. "Resultatene fra kapittel 6 viser at" → kapittel 7
  2. Fjern ledende mellomrom i Konklusjon-H2-overskrifter
  3. Tabell 3.1 → 4.1, Tabell 5.1 → 6.1
  4. Figur 6.1–6.4 → 7.1–7.4, Figur 6.5 → 8.1
  5. Legg til tabellkaption for tabeller uten tittel
  6. Endre "Kodebase, data og referanser" fra H1 til H2
"""
from docx import Document
from docx.oxml import OxmlElement

DOC = '/Users/oskareide/Desktop/LOG650 - nyeste - ren.docx'
NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc  = Document(DOC)
body = doc.element.body
para_map = {id(p._element): p for p in doc.paragraphs}
fixes = []


# ── helpers ───────────────────────────────────────────────────────────────────

def get_text(elem):
    return ''.join(t.text or '' for t in elem.iter(f'{{{NS}}}t'))

def replace_in_para(p, old, new):
    if old not in p.text:
        return False
    full = p.text.replace(old, new)
    runs = p.runs
    if not runs:
        return False
    runs[0].text = full
    for r in runs[1:]:
        r.text = ''
    return True

def make_caption(text):
    """Paragraph with table-caption text (Normal style)."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(f'{{{NS}}}val', 'Normal')
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i = OxmlElement('w:i')   # italics for caption
    rPr.append(i)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


# ── 1–4: paragraph-level text replacements ────────────────────────────────────

for p in doc.paragraphs:
    t = p.text

    # 1. Diskusjon – "Resultatene fra kapittel 6"
    if 'Resultatene fra kapittel 6 viser' in t:
        if replace_in_para(p, 'Resultatene fra kapittel 6 viser',
                           'Resultatene fra kapittel 7 viser'):
            fixes.append('"Resultatene fra kapittel 6 viser" → kapittel 7')

    # 2. Leading spaces in Konklusjon H2 headings
    if p.style.name == 'Heading 2' and t.startswith(' '):
        runs = p.runs
        if runs:
            runs[0].text = runs[0].text.lstrip()
            fixes.append(f'Fjernet ledende mellomrom i H2: "{t.strip()[:40]}"')

    # 3. Tabellnumre
    if 'Tabell 3.1' in t:
        if replace_in_para(p, 'Tabell 3.1', 'Tabell 4.1'):
            fixes.append('"Tabell 3.1" → "Tabell 4.1"')
    if 'Tabell 5.1' in t:
        if replace_in_para(p, 'Tabell 5.1', 'Tabell 6.1'):
            fixes.append('"Tabell 5.1" → "Tabell 6.1"')

    # 4. Figurnumre
    fig_map = {
        'Figur 6.1': 'Figur 7.1',
        'Figur 6.2': 'Figur 7.2',
        'Figur 6.3': 'Figur 7.3',
        'Figur 6.4': 'Figur 7.4',
        'Figur 6.5': 'Figur 8.1',
    }
    for old_fig, new_fig in fig_map.items():
        if old_fig in t:
            if replace_in_para(p, old_fig, new_fig):
                fixes.append(f'"{old_fig}" → "{new_fig}"')
                break


# ── 5. Legg til tabellkapsler for tabeller uten tittel ────────────────────────

# Find tables by context of previous sibling (robust: uses lxml getprevious/getnext)
tbl_contexts = {
    'Deskriptiv analyse':
        'Tabell 7.1 – Deskriptiv statistikk for de fire produktene',
    'Stasjonaritetstest (ADF)':
        'Tabell 7.2 – ADF-testresultater (log-differensierte serier)',
    'SARIMA-parameterestimater':
        'Tabell 7.3 – SARIMA-parameterestimater og modelltilpasning (AIC, BIC)',
    'Testperiode: ca. november':
        'Tabell 7.4 – Modellsammenlikning (RMSE, MAE, MAPE) for testperioden',
    'Beregnet med z = 1,645':
        'Tabell 8.1 – Beregnede lagerstyringsstørrelser (95 % servicegrad)',
    'Scenario A: erfaringsbasert':
        'Tabell 8.2 – Simuleringsresultater: enheter med utsalg per scenario',
}

for child in list(body):
    tag = child.tag.split('}')[-1]
    if tag != 'tbl':
        continue
    prev = child.getprevious()   # lxml sibling – always correct
    if prev is None:
        continue
    prev_text = get_text(prev).strip()
    for key, caption in tbl_contexts.items():
        if prev_text.startswith(key):
            nxt = child.getnext()
            if nxt is not None and caption[:20] in get_text(nxt):
                break  # already inserted
            cap_elem = make_caption(caption)
            child.addnext(cap_elem)
            fixes.append(f'Lagt til: {caption}')
            break


# ── 6. "Kodebase, data og referanser" – H1 → H2 ──────────────────────────────

for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and p.text.strip() == 'Kodebase, data og referanser':
        pStyle = p._element.find(f'.//{{{NS}}}pStyle')
        if pStyle is not None:
            pStyle.set(f'{{{NS}}}val', 'Overskrift2')
            fixes.append('"Kodebase, data og referanser" endret fra H1 til H2')
        break


# ── lagre ─────────────────────────────────────────────────────────────────────

doc.save(DOC)
print(f'\n{len(fixes)} rettelse(r):')
for f in fixes:
    print(f'  ✓ {f}')

#!/usr/bin/env python3
"""
Moves content from LOG650 - BM.docx into the template structure (Mal prosjekt LOG650 v2.docx).
Output: LOG650 - BM - Mal.docx on the Desktop.

Chapter mapping:
  Sammendrag            → Sammendrag
  1 Introduksjon        → Innledning (1.1 Problemstilling, 1.3 Avgrensninger, rest kept)
  2 Teori og litteratur → 2.0 Litteratur (intro only) + 3.0 Teori (2.1-2.4)
  3 Casebeskrivelse     → 4.0 Casebeskrivelse
  4 Data og metode      → 5.0 Metode og data
  5 Modellering         → 6.0 Modellering
  6 Analyse og res.     → 7.0 Analyse (6.1-6.5) + 8.0 Resultat (6.6-6.7)
  7 Diskusjon           → 9.0 Diskusjon
  8 Konklusjon          → 10.0 Konklusjon
  Referanser            → 11.0 Bibliografi
                        → 12.0 Vedlegg (empty)
"""
import re
import copy
from docx import Document
from docx.oxml import OxmlElement

ASN_PATH  = '/Users/oskareide/Desktop/LOG650 - BM.docx'
TMPL_PATH = '/Users/oskareide/Desktop/Mal prosjekt LOG650 v2.docx'
OUT_PATH  = '/Users/oskareide/Desktop/LOG650 - BM - Mal.docx'


# ── helpers ──────────────────────────────────────────────────────────────────

def strip_num(text):
    """Remove leading chapter numbers: '2.1 Lagerstyring' → 'Lagerstyring'"""
    return re.sub(r'^\d+(\.\d+)*\s+', '', text)

HEADING_STYLES = {1: 'Overskrift1', 2: 'Overskrift2', 3: 'Overskrift3'}

def make_heading(text, level=1):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', HEADING_STYLES[level])
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_h1(text):
    return make_heading(text, level=1)

def replace_run_text(p_elem, new_text):
    """Overwrite all runs in a paragraph with a single run containing new_text."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r in p_elem.findall(f'{{{ns}}}r'):
        p_elem.remove(r)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = new_text
    r.append(t)
    p_elem.append(r)


# ── load documents ────────────────────────────────────────────────────────────

asn  = Document(ASN_PATH)
tmpl = Document(TMPL_PATH)
body = tmpl.element.body


# ── group assignment content by H1 section ───────────────────────────────────

sections = {}
cur = 'FRONT'
sections[cur] = []

for p in asn.paragraphs:
    if p.style.name == 'Heading 1' and p.text.strip():
        cur = p.text.strip()
        sections[cur] = []
    else:
        sections.setdefault(cur, [])

# Re-collect as raw XML elements (paragraphs AND tables, in order)
raw_sections = {k: [] for k in sections}
cur = 'FRONT'
for child in asn.element.body:
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        from docx.oxml.ns import qn
        style_elem = child.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
        style_val  = style_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') \
                     if style_elem is not None else ''
        text = ''.join(t.text or '' for t in
                       child.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')).strip()
        if style_val in ('Overskrift1',) and text:
            cur = text
            raw_sections.setdefault(cur, [])
            continue
    if tag in ('p', 'tbl'):
        raw_sections.setdefault(cur, []).append(child)

# Also check style via python-docx name (some docs use style name not id)
# Rebuild using paragraph API to be safe
raw_sections2 = {}
cur = 'FRONT'
raw_sections2[cur] = []
para_iter = iter(asn.paragraphs)
elem_iter  = list(asn.element.body)

# Use a combined approach: iterate body children, check heading via paragraph style
para_map = {id(p._element): p for p in asn.paragraphs}
cur = 'FRONT'
raw_sections2[cur] = []
for child in asn.element.body:
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        para = para_map.get(id(child))
        if para and para.style.name == 'Heading 1' and para.text.strip():
            cur = para.text.strip()
            raw_sections2.setdefault(cur, [])
            continue
        raw_sections2.setdefault(cur, []).append(child)
    elif tag == 'tbl':
        raw_sections2.setdefault(cur, []).append(child)

raw_sections = raw_sections2
print("Sections:", [k for k in raw_sections if k != 'FRONT'])


# ── cut template at first H1 ("Innledning") ──────────────────────────────────

children = list(body)
cut_at = None
for i, child in enumerate(children):
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        para = None
        for p in tmpl.paragraphs:
            if p._element is child:
                para = p
                break
        if para and para.style.name == 'Heading 1':
            cut_at = i
            print(f"Cutting template at index {i}: '{para.text}'")
            break

if cut_at is None:
    raise RuntimeError("Could not find first Heading 1 in template")

for child in children[cut_at:]:
    if child.tag.split('}')[-1] != 'sectPr':
        body.remove(child)


# ── append helper ─────────────────────────────────────────────────────────────

def is_heading(elem, levels=(2, 3, 4)):
    """Check if elem is a heading at one of the given levels using para_map."""
    para = para_map.get(id(elem))
    if para is None:
        return False
    return any(para.style.name == f'Heading {lvl}' for lvl in levels)

def add(elems, strip_numbers=True):
    """Append deep-copied elements to body. Strip chapter numbers from headings."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for child in elems:
        new = copy.deepcopy(child)
        tag = new.tag.split('}')[-1]
        if tag == 'p' and strip_numbers and is_heading(child, levels=(2, 3, 4)):
            text = ''.join(t.text or '' for t in new.iter(f'{{{ns}}}t')).strip()
            clean = strip_num(text)
            if clean != text:
                replace_run_text(new, clean)
        body.append(new)


# ── build new document ────────────────────────────────────────────────────────

# Sammendrag
body.append(make_h1('Sammendrag'))
add(raw_sections.get('Sammendrag', []))

# 1.0 Innledning
body.append(make_h1('Innledning'))
add(raw_sections.get('1 Introduksjon', []))

# 2.0 Litteratur  (intro paragraph of kap 2, before first H2)
body.append(make_h1('Litteratur'))
teori = raw_sections.get('2 Teori og litteratur', [])
for child in teori:
    if is_heading(child, levels=(2, 3)):
        break
    body.append(copy.deepcopy(child))

# 3.0 Teori  (2.1–2.4)
body.append(make_h1('Teori'))
after_first_h2 = False
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for child in teori:
    if is_heading(child, levels=(2, 3)):
        after_first_h2 = True
    if after_first_h2:
        new = copy.deepcopy(child)
        if is_heading(child, levels=(2, 3)):
            text = ''.join(t.text or '' for t in new.iter(f'{{{ns}}}t')).strip()
            clean = strip_num(text)
            if clean != text:
                replace_run_text(new, clean)
        body.append(new)

# 4.0 Casebeskrivelse
body.append(make_h1('Casebeskrivelse'))
add(raw_sections.get('3 Casebeskrivelse', []))

# 5.0 Metode og data
body.append(make_h1('Metode og data'))
add(raw_sections.get('4 Data og metode', []))

# 6.0 Modellering
body.append(make_h1('Modellering'))
add(raw_sections.get('5 Modellering', []))

# 7.0 Analyse (6.1–6.5) and 8.0 Resultat (6.6–6.7)
analyse_all = raw_sections.get('6 Analyse og resultater', [])
split_idx = None
for i, child in enumerate(analyse_all):
    if is_heading(child, levels=(2, 3)):
        text = ''.join(t.text or '' for t in child.iter(f'{{{ns}}}t')).strip()
        if '6.6' in text or 'Lagerstyringsstørrelser' in text:
            split_idx = i
            break

analyse_part  = analyse_all[:split_idx] if split_idx else analyse_all
resultat_part = analyse_all[split_idx:] if split_idx else []

body.append(make_h1('Analyse'))
add(analyse_part)

body.append(make_h1('Resultat'))
add(resultat_part)

# 9.0 Diskusjon
body.append(make_h1('Diskusjon'))
add(raw_sections.get('7 Diskusjon', []))

# 10.0 Konklusjon
body.append(make_h1('Konklusjon'))
add(raw_sections.get('8 Konklusjon', []))

# 11.0 Bibliografi
body.append(make_h1('Bibliografi'))
add(raw_sections.get('Referanser', []), strip_numbers=False)

# 12.0 Vedlegg
body.append(make_h1('Vedlegg'))

# ── save ──────────────────────────────────────────────────────────────────────
tmpl.save(OUT_PATH)
print(f"\nSaved → {OUT_PATH}")

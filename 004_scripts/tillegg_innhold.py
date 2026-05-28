#!/usr/bin/env python3
"""
Tre innholdstillegg i LOG650 - nyeste - ren.docx:
  1. Sammendrag – konkrete tall fra Tabell 8.2
  2. Teori, Analyse, Resultat – innledningsavsnitt
  3. Litteratur-kapittelet – faktisk litteraturgjennomgang
"""
from docx import Document
from docx.oxml import OxmlElement

DOC = '/Users/oskareide/Desktop/LOG650 - nyeste - ren.docx'
NS  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc  = Document(DOC)
body = doc.element.body
para_map = {id(p._element): p for p in doc.paragraphs}
fixes = []


# ── helpers ───────────────────────────────────────────────────────────────────

def replace_in_para(p, old, new):
    if old not in p.text:
        return False
    full = p.text.replace(old, new)
    runs = p.runs
    if not runs:
        return False
    runs[0].text = full
    for r in runs[1:]:
        r.text = ''
    return True

def make_para(text, style_val='Normal'):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(f'{{{NS}}}val', style_val)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


# ── 1. Sammendrag – legg til konkrete tall ────────────────────────────────────

for p in doc.paragraphs:
    if 'En simulering viser at et KI-støttet innkjøpssystem' in p.text \
            and 'for alle fire produkter' in p.text \
            and '15 307' not in p.text:
        old = 'for alle fire produkter.'
        new = ('for alle fire produkter. '
               'For terrassebord (28x120 mm) reduseres enheter med utsalg fra 15 307 til 1 410, '
               'og for konstruksjonstre (48x98 mm) fra 16 616 til 1 384. '
               'For festemidlene er utsalg tilnærmet eliminert.')
        if replace_in_para(p, old, new):
            fixes.append('Sammendrag – konkrete simuleringstall lagt til')
        break


# ── 2. Innledningsavsnitt til Teori, Analyse og Resultat ─────────────────────

intro_texts = {
    'Teori': (
        'Dette kapittelet presenterer det teoretiske grunnlaget som prosjektet bygger på. '
        'De fire temaene — lagerstyring, etterspørselsprognoser, kunstig intelligens i logistikk '
        'og KI-støttede innkjøpssystemer — danner rammeverket for modellering, analyse og '
        'diskusjon i de påfølgende kapitlene.'
    ),
    'Analyse': (
        'Dette kapittelet presenterer analysen av de historiske salgsdata for de fire produktene. '
        'Analysen følger en femstegsstruktur: fra deskriptiv statistikk og stasjonaritetstest '
        'via ACF/PACF-analyse og parameterestimering til modellsammenlikning på testsettet.'
    ),
    'Resultat': (
        'Dette kapittelet presenterer de beregnede lagerstyringsstørrelsene og resultatene fra '
        'lagersimuleringen. Beregningene bygger på prognoseresultatene fra kapittel 7 og '
        'lagerstyringsformler presentert i kapittel 6.'
    ),
}

for child in list(body):
    tag = child.tag.split('}')[-1]
    if tag != 'p':
        continue
    para = para_map.get(id(child))
    if para is None:
        continue
    if para.style.name == 'Heading 1' and para.text.strip() in intro_texts:
        ch_name = para.text.strip()
        # Check that next sibling is NOT already a Normal paragraph intro
        nxt = child.getnext()
        nxt_para = para_map.get(id(nxt)) if nxt is not None else None
        if nxt_para and nxt_para.style.name == 'Normal' \
                and len(nxt_para.text.strip()) > 30:
            continue  # already has intro
        intro_elem = make_para(intro_texts[ch_name])
        child.addnext(intro_elem)
        fixes.append(f'Innledningsavsnitt lagt til i kapittel: {ch_name}')


# ── 3. Litteratur-kapittelet – full litteraturgjennomgang ─────────────────────

litt_intro_marker = 'Dette kapittelet gir en oversikt over sentrale vitenskapelige arbeider'

litt_paragraphs = [
    ('Lagerstyring og forsyningskjedeteori',  'h2'),
    ('Lagerstyring er et veletablert forskningsfelt med røtter tilbake til Harris (1913) sin '
     'klassiske EOQ-modell, som definerer den optimale bestillingsmengden som minimerer summen '
     'av bestillings- og lagerkostnader. Silver, Pyke og Thomas (2017) har videreutviklet dette '
     'rammeverket for moderne forsyningskjeder og viser hvordan sikkerhetslager og bestillingspunkt '
     'kan dimensjoneres under etterspørselsusikkerhet. Chopra og Meindl (2016) understreker at '
     'valget av lagerstyringssystem — kontinuerlig versus periodisk gjennomgang — er tett knyttet '
     'til ERP-systemets kapabiliteter og leverandørenes fleksibilitet.',
     'normal'),
    ('Etterspørselsprognoser',                'h2'),
    ('Hyndman og Athanasopoulos (2021) gir en systematisk gjennomgang av kvantitative '
     'prognosemetoder og viser at SARIMA-modeller er særlig velegnet for sesongbaserte tidsserier '
     'med begrenset datahistorikk. Makridakis, Spiliotis og Assimakopoulos (2018) sammenligner '
     'statistiske og maskinlæringsbaserte metoder i en stor empirisk studie og konkluderer med '
     'at ingen enkeltmodell dominerer på tvers av alle datasett — valg av metode bør tilpasses '
     'datakarakteristika og prognosehorisonten.',
     'normal'),
    ('Kunstig intelligens i logistikk',       'h2'),
    ('Carbonneau, Laframboise og Vahidov (2008) demonstrerer at maskinlæringsmodeller kan gi '
     'høy prognoseakkuratesse for etterspørselsdata i varehandel, særlig ved ikke-lineære '
     'etterspørselsmønstre. Toorajipour mfl. (2021) gjennomgår systematisk forskningen på KI '
     'i supply chain management og peker på etterspørselsplanlegging og beslutningsstøtte som '
     'de to områdene med sterkest dokumentert effekt. Fosso Wamba mfl. (2015) finner at '
     'datakvalitet og organisatorisk tilpasning er de viktigste flaskehalsene ved implementering '
     'av KI-løsninger i eksisterende systemer.',
     'normal'),
    ('Forskningsgap',                         'h2'),
    ('Felles for mye av litteraturen er at implementeringsstudier retter seg mot store bedrifter '
     'med rik datahistorikk og dedikerte dataressurser. Det finnes begrenset forskning på hvordan '
     'disse metodene kan tilpasses mellomstore varehandelsaktører med korte tidsserier og '
     'begrensede IT-ressurser. van der Vorst, Beulens og van Beek (2000) peker på at '
     'informasjonsflyt og systemintegrasjon er avgjørende for vellykkede implementeringer også '
     'i enklere forsyningskjedesettinger. Dette prosjektet bidrar til å fylle dette gapet ved '
     'å anvende etablerte metoder i en konkret SMB-kontekst.',
     'normal'),
]

# Find the Litteratur intro paragraph and check if H2 sections already exist
for child in list(body):
    tag = child.tag.split('}')[-1]
    if tag != 'p':
        continue
    para = para_map.get(id(child))
    if para is None:
        continue
    if litt_intro_marker in para.text:
        # Check if H2 already added
        nxt = child.getnext()
        nxt_para = para_map.get(id(nxt)) if nxt is not None else None
        if nxt_para and nxt_para.style.name in ('Heading 2', 'Overskrift2'):
            fixes.append('Litteratur-H2 allerede på plass – ingenting lagt til')
            break
        # Insert all paragraphs after the intro
        ref = child
        for text, kind in litt_paragraphs:
            if kind == 'h2':
                new_elem = make_para(text, 'Overskrift2')
            else:
                new_elem = make_para(text, 'Normal')
            ref.addnext(new_elem)
            ref = new_elem
            para_map[id(new_elem)] = None  # avoid re-matching
        fixes.append('Litteratur-kapittelet – fire seksjoner med gjennomgang lagt til')
        break


# ── lagre ─────────────────────────────────────────────────────────────────────

doc.save(DOC)
print(f'\n{len(fixes)} tillegg lagret:')
for f in fixes:
    print(f'  ✓ {f}')
